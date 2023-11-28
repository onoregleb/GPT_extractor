import streamlit as st
from langchain.chains import RetrievalQA
import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain.chat_models import ChatOpenAI
from langchain.document_loaders import TextLoader
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from docx import Document
import asyncio
import nest_asyncio
import io
from openai.error import APIError
import pysqlite3
import sys
nest_asyncio.apply()


sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')


class MyDocument:
    def __init__(self, page_content, metadata=None):
        self.page_content = page_content
        if metadata is None:
            metadata = {"source": "Generated by ChatGPT"}
        self.metadata = metadata


glob = """You are a useful assistant with artificial intelligence for creating business plans for large texts.
Use the following context snippets to answer the question at the end (Question).
This answer should contain the answer to the question to the answer, as well as contain additional information in order to be
as capacious as possible and at the same time relevant.
If you don't know the answer, just say you don't know, don't try to come up with an answer.
First make sure that the attached text is relevant to the question asked.
If the text is irrelevant to the question, just say that the text is irrelevant.
If it seems to you that the answer can be supplemented with information from the Internet to make the text more understandable
and interesting for investors, supplement it. However, first check whether the answer fits the meaning of the question.
Use a maximum of fifteen sentences. Give the answer as detailed as possible. Give me an answer in Russian.
Context: {context}
Question: {question}
Useful answers:"""

os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]

response = Document()


def docx_bytesio_to_md(input_bytesio):
    try:
        doc = Document(input_bytesio)
        md_content = ""
        for paragraph in doc.paragraphs:
            text = paragraph.text
            md_content += text + '\n'
        with open("output_file.txt", 'w', encoding='utf-16') as md_file:
            md_file.write(md_content)
    except Exception as e:
        print(f'Произошла ошибка при конвертации и сохранении: {e}')



def generate_response(uploaded_file, employees_info, project_name, num_employees):
    if uploaded_file is not None:
        question1 = f"""Представь, что геймдев разработчик и у тебя есть игра, где сюжет разворачивается вокруг проекта {project_name}, опиши его суть. Твоя задача создать для этого проекта {num_employees} сотрудника(-ов) и указанными именами и данными из ВУЗа и 
        придумать им небольшой бэкграунд, т.е. опыт работы, присвоить должность и описать какие обязанности будет исполнять на работе, таким образом, чтобы они хорошо подходили под проект, но были не супер подходящими кандидатами, к примеру т.к. это студенты, то опыт у них маленький. 
        Когда будешь говорить про опыт, то пусть они имеют уже какие-то небольшие заслуги, к примеру, программист на звании junior или сварщик с разрядом, также придумай, что они работают уже какое-то количество лет и делали такой-то кейс. 
        Входные данные: [Название проекта] - {project_name}. ([Суть проекта]) - опиши кратко, но емко. [Количество сотрудников] -{num_employees} ([ФИО 1]){employees_info[0]['employee_name']} ([ВУЗ, факультет и курс сотрудника 1]) {employees_info[0]['university_info']} и так далее для всех {employees_info}.
        Формат выходных данных: 
        1. Первый студент 
        1.1. ([ФИО]) обучающийся в (таком то вузе, на такой то специальности, на таком-то курсе ([ВУЗ, факультет и курс сотрудника 3])) будет занимать (такую-то должность) в проекте. 
        1.2. В этом пункте воссоздаёшь его бэкграунд, т.е. опыт предыдущей работы, участия в конкурсах и прочие заслуги, которые оцениваются при приёме на работу. Объем текста пункта должен быть от 1000 до 2000 символов. Добавь немного конкретики 
        1.3. Написать какие обязанности будет исполнять Для других сотрудников структура и подход аналогичный. """
        question2 = """Представь, что ты опытный маркетолог и твоя задача написать цепляющую аннотацию для продукта, который только собираются выводить на рынок.
        Информацию для аннотации возьми из файла. Представь, что ты только в начале своего пути разработчика и ты создаёшь эту аннотацию для бизнес-ангела,
        а сам только хочешь открыть компанию и заниматься этим делом. Ответ должен не превышать по объёму 1000 символов и выведен в следующем формате:
        1. Описание продукта с подчеркиванием научной составляющей проекта
        2. Цель продукта с ориентиром на конечного пользователя
        3. Задачи которые необходимо выполнить для достижения цели
        4. Ожидаемые результаты через 1-2 месяца и через год, в случае если инвестор выделит тебе миллион рублей на проект
        5. Кто будет целевой аудиторией для проекта
        Дай ответ на русском. Пиши ответ не по пунктам, а сплошным текстом."""
        question3 = """"Представь, что ты маркетолог и тебя наняли в проект, который только собирает открываться. К тебе пришёл шеф и говорит написать текст
        для знакомства инвестора с продуктом. В файле, который я прислал, содержится информация о продукте, используй её при описании продукта. Шеф говорит тебе,
        что после прочтения раздела у инвестора должно сложиться понимание сути бизнеса и потребностей клиентов, которые бизнес удовлетворяет, а также – понимание УТП,
        конкурентных преимуществ.
        В этом разделе описывается тот продукт или услуга, на производство которого и направлена наша деятельность. На данном этапе нужно постараться воссоздать
        образ и детали будущей продукции или услуги, с помощью чего планируется получать прибыль.
        Таким образом, можно выделить 5 пунктов, которые необходимо раскрыть в разделе:
        1. что планируется производить, какой товар или услуга;
        2. качественное описание общих и уникальных свойств;
        3. состав, клиентоориентированность и направленность;
        4. географический охват;
        5. сегмент покупателей.
        Дай ответ на русском. Пиши ответ не по пунктам, а сплошным текстом."""
        question4 = """Представь, что ты маркетолог и тебя наняли в проект, который только собирает открываться. К тебе пришёл шеф и говорит написать
        описание целевой аудитории для представления инвестору. В файле, который я прислал, содержится информация о продукте, используй её. В описании
        тебе нужно сказать о том:
        - ЦА это юридические и (или) физические лица?;
        - для юридических лиц указать категорию бизнеса (крупный/средний/малый бизнес,
        отрасль и проч.), обосновать выбор категории и продемонстрировать примеры сфер этих компаний;
        - для физических лиц дать описание их характеристик, согласно которым у них
        возникает потребность в приобретении производимого предприятием продукта
        или услуги;
        - указать географическое расположение потенциальных потребителей и проч.;
        - на какой сектор рынка (B2B/ B2C) ориентируется проект;
        - выявить проблемы и желания клиентов;
        - предположить какие характеристики продукта будут ключевыми и обосновать выбор.
        Свой ответ пиши в формате отчёта. Если информации из файла недостаточно, то ответь на вопросы, используя свой опыт в маркетинге."""
        question5 = """Представь, что ты инженер и тебя наняли в проект, который только собирает открываться. К тебе пришёл шеф и говорит описать научно-технические
        решения, которые используются в процессе разработки продукта, чтобы представить это инвестору. В файле, который я прислал, содержится информация о продукте,
        используй её. Шеф говорит тебе, чтобы ты наметил решения, которые будут использованы в процессе. Сюда входят: программы, среды программирования, языки
        программирования, техника, алгоритмы и т.д., т.е. все те вещи, инструменты, алгоритмы и идеи, объединив которые мы получим продукт.
        Свой ответ пиши в формате отчёта. Если информации из файла недостаточно, то ответь на вопросы, используя свой опыт в этой сфере. """

        question6 = """Представь, что ты генеральный директор новой компании, которая только открылась и собирается выводить на рынок новый продукт. Информация о нём содержится в файле, который я тебе прислал. Твоя задача - создать корпоративный документ, в котором будут отражены следующие пункты:
        - описание принципов/алгоритмов того, как будет организован бизнес;
        - как будут выстраиваться взаимоотношения между создаваемым юридическим лицом (т.е.
        предприятием) и предполагаемыми потребителями и поставщикам;
        - будет ли организовываться сеть посредников;
        - планируется ли участие в каком-либо агрегаторе, акселераторе, стартап-академии.
        Т.к. этот документ периодически будет показываться инвесторам, то ты должен понимать, что тебе необходимо положительно раскрыть эту тему, дав
        положительный задел твоей компании, т.е. и поучаствовать в каких-то акселераторах.

        Прими к сведению следующую информацию и в ответе на вопросе используй это не как базис размышления, а как дополнительную информацию, которая дополнит ответ.
        Компания – сложная система, которая в свою очередь состоит других сложных подсистем. Сделано это для того, чтобы чётко разграничить зоны и получить
        максимальный результат от каждой зоны и специалиста. Поэтому организацию стартапа можно рассматривать с нескольких точек зрения:
        1. Зона продукта. Это направление занимается созданием продукта, добавлением возможностей, исправлением косяков, сбором и анализом информации для
        последующих модернизаций – в общем делает так, чтобы потребитель был доволен. Тут мы говорим о том, как будет функционировать продукт, т.е. что нужно
        сделать, чтобы работало. К примеру, взять А, положить в Б, добавить В, получить из этого всего наше Г.
        2. Зона роста. Весь бизнес должен развиваться. Для этого требуется где-то привлечение инвесторов, где-то сотрудничество и партнёрство, прохождение
        акселерационных программ для оттачивания KPI и показателей, повышение квалификации команды, участие на научных форумах для рекламы. Всестороннее
        развитие компании помогает перейти на новый уровень.
        3. Зона «бумажных шестерёнок». Расписываем процесс регистрации компании, с выделением особенностей компании (ОБЯЗАТЕЛЬНО это ООО, усн 6%).
        Назначить на должности сотрудников и в двух словах рассказать о том, чем они будут заниматься."""

        question7 = """Представь, что ты продуктовый аналитик и по совместительству маркетолог новой компании, которая только собирается открываться и собирается
        выводить на рынок новый продукт. Информация о нём содержится в файле, который я тебе прислал. Твоя задача - привести обоснования реализуемости бизнеса,
        а именно показать, что проект может быть реализован, быть успешным, приносить доход, занять часть рынка и т.д. Для этого необходимо обязательно ответить на эти вопросы:
        - Какую пользу для общества и потенциальной ЦА несёт продукт?
        - Насколько конкурентоспособным будет новый товар или услуга и почему?
        - В чем заключается новизна предлагаемой продукции?
        - Почему идея потенциальна прибыльна?

        И ответить на некоторые вопросы из этого списка, которые с положительной точки будут аргументами к теме реализации проекта:
        - Существуют ли на рынке аналоги или сходные товары?
        - На какую аудиторию ориентирован товар? В чем ее особенности?
        - Какие каналы сбыта?
        - Есть ли у вас опыт создания подобных проектов?
        - Есть ли поддержка со стороны (партнеров, инвесторов, госорганов и т.д.)?
        - Есть ли у вас специалисты, компетентные в данном виде деятельности? (Тут надо ответить да, что есть команда, но не профессионалы,
        а обычные студенты, которым интересно это направление и уже сформирована команда, но другими словами)
        - Какой объём капиталовложений понадобится для реализации бизнес-идеи?
        - Сколько вы можете вложить без привлечения инвесторов?
        - Требует ли данный вид деятельности получения лицензии или сертификата?
        - Если да, то насколько сложным и длительным будет процесс подготовки разрешительной документации?
        - Есть ли помещение или земельный участок, оборудование?
        - Насколько длительным будет срок окупаемости?
        - Какие существуют перспективы развития бизнеса?"""

        question8 = """ Представь, что ты главный специалист в компании по её продукту в новой компании, которая только собирается открываться и собирается
        выводить на рынок новый продукт. Информация о продукте содержится в файле, который я тебе прислал. Твоя задача - создать документ, в котором будут отражены
        технические параметры разрабатываемого нами продукта, информация о котором имеется у тебя а именно:
        - 5-7 цифровых показателей в цифрах, которые характеризуют разработку продукта (к примеру, если мы производим тракторы, то сколько 1 трактор сможет
        обработать гектаров пшеницы за 1 единицу времени, т.е. что-то что отражает работу продукта);
        - визуализация: архитектура разработки, схема функционирования, дизайн продукта;
        - входные и выходные воздействия (к примеру, если продукт трактор, то к входным воздействиям можно отнести топливо, которым он заправляется,
        человека, который будет управлять трактором, выходные, соответственно, то, что остаётся на выходе);
        - условия эксплуатации продукта (к примеру, если это какой-то онлайн сервис, то пользователю понадобится компьютер, с установленным браузером и
        выходом в интернет и т.д.).
        Если недостаточно информации из файла, используй свой опыт и знания, чтобы раскрыть этот раздел."""

        question9 = """ Представь, что ты Chief Operating Officer в новой компании, которая только собирается открываться и
        собирается выводить на рынок новый продукт. Информация о продукте содержится в файле, который я тебе прислал. Твоя задача - написать текст,
        в котором будет описана бизнес модель Остервальдера-Пинье с её 9 пунктами для проекта в прогнозе на ближайший год, с учётом того, что ты можешь получить
        финансирование от инвестора в размере миллиона рублей на основании информации из файла.
        Если недостаточно информации из файла, используй свой опыт и знания, чтобы раскрыть этот раздел. Для поиска актуальной и недостоюзей информации используй "Search" """

        question10 = """Представь, что ты маркетолог в новой компании, которая только собирается открываться и собирается выводить на рынок новый продукт.
        Информация о продукте содержится в файле, который я тебе прислал. Твоя задача - создать документ, в котором будут описаны конкурентные преимущества
        нашего продукта относительно ближайших конкурентов. Ответ давай в следующем формате:
        - распиши 10 УТП (уникальных торговых предложений) нашего продукта;
        - потом найди 6 конкурентов, дай описание их продукту (укажи страну), и к каждому конкуренту приведи 5 сравнений по технико-экономическим параметрам,
        по которым наш продукт уже лучше или может быть лучше потенциально. Конкуренты могут быть как прямые, так и косвенные, как отечественные (т.е. из России), так и заграничные.
        Пиши ответ от лица разработчиков. Если недостаточно информации из файла, используй свой опыт и знания, чтобы раскрыть этот раздел."""

        question11 = """Представь, что ты главный специалист в компании по её продукту в новой компании, которая только собирается открываться и собирается
        выводить на рынок новый продукт. Информация о продукте содержится в файле, который я тебе прислал. Твоя задача - написать текст, в котором будут
        отражено научно-техническое решение для создания продукта, к примеру, можно использовать эти вопросы:
        - как будет разрабатываться метод или алгоритм;
        - какой научный подход или идея лежит в основе разработки;
        - какие математические модели будут применяться;
        - какой язык программирования будет использоваться;
        - какие материалы, компоненты будут использоваться;
        - на каких данных будет обучаться ИИ;
        - и т.д.
        Составь структуру, опиши функционал функционала, предоставь примерные схемы, пропиши нюансы, которые важны, чтобы команда смогла не
        упустить их при разработке, в том числе нужно прояснить пожелания к дизайну. Обозначь возможный технологический стек, и используйте наиболее
        подходящие для вашего проекта решения. Распиши сценарии использования продукта.
        Сценарий нужен для понимания принципа работы продукта. Например, если область работы касается IT, сценарий отвечает на вопрос «Как будет вести
        себя пользователь?» и дает понимание главных функций сайта. Соблюдай установленные стандарты, к числу которых относятся нормативные документы,
        такие как ГОСТы или ISO.
        Если недостаточно информации из файла, используй свой опыт и знания, чтобы раскрыть этот раздел."""

        question12 = """Представь, что ты главный специалист в компании по её продукту в новой компании, которая только собирается открываться
        и собирается выводить на рынок новый продукт. Информация о продукте содержится в файле, который я тебе прислал.
        Твоя задача - написать текст, в котором будет отражен задел научно-технического решения на данный момент, чтобы инвестор выдал тебе деньги.
        Но, пообщавшись с командой, вы пришли к мнению, что вы не будете раскрывать все карты, т.е. не будете говорить о том, что у вас уже
        выполнена разработка по тому проекту, который я прислал файлом, а скажете, что вы только собирайтесь разрабатывать этот проект.
        К примеру, к заделу можно будет отнести компьютеры, команду, которая разбирается в этом направлении, программы, которые будут
        использоваться, материалы, алгоритмы, изученные методы, модели, схемы, лицензии, патенты и т.д. Но стоит говорить об этом аккуратно, чтобы потом в конце,
        информация из файла оказалась конечной точкой продукта, а сейчас вы как будто ещё собирайтесь приступать за этот проект.
        Если недостаточно информации из файла, используй свой опыт и знания, чтобы раскрыть этот раздел. Напиши ответ в качестве отчета"""

        question13 = """Представь, что ты главный маркетолог продукта в новой компании, которая только собирается открываться и собирается выводить на рынок новый
        продукт. Информация о продукте содержится в файле, который я тебе прислал. Твоя задача - создать документ, в котором будут отражены боли, проблемы,
        желания и потребности целевой аудитории.
        Пример:
        Есть фермер, который живёт за городом, раз в неделю он идёт на рынок продавать свои овощи и фрукты. Но у него есть
        ПРОБЛЕМА (Тут начинается описание проблемы) – чтобы попасть в город, ему нужно пройти по реке, где он мочит ноги, одежду, обувь и часть товара,
        зимой вообще идти опасно, сколько упал, сломал ногу/руку – кирдык башка. Из-за этой реки он уже несколько раз заболел, намочил свои деньги, в связи
        с чем он не мог дать нормально сдачу и испортил внешний вид товара. Каким образом он решает проблему? Он решил построить мостик через реку. Для этого ему
        понадобится срубить 40 деревьев, подготовить дерево, покрыть лаком, найти верёвки, связать, построить, вбить и т.д., чтобы в итоге получился мост. Он это
        делает – решает свою проблему.
        Необходимо расписать не только проблемы, но и боли (т.е. то, что серьёзнее проблемы), желания и потребности, на сколько сильно это волнует нашу
        целевую аудиторию (предварительно сформулируй понятие о целевой аудитории проекта в соотсветствии с общепринятыми критериями и градациями),
        как аудитория самостоятельно пытается справиться с этим/закрывает. А так же распределить их среди 3 основных сегментов целевой аудитории.
        Формат ответа:
        1. Боли:
        1.1. ...
        1.2. ...
        и т.д.
        2. Проблемы
        2.1. ...
        2.2. ...
        и т.д.
        3. Потребности
        3.1. ...
        3.2. ...
        и т.д.
        4. Желания
        4.1. ...
        4.2. ...
        и т.д.
        5. Сегменты
        5.1. к первому сегменту с (описание сегмента) можно отнести такие-то пункты (к примеру, 1.1., 3.4 и т.д.)
        5.2. аналогично 5.1. только своей целевой аудитории. В ответе старайся оперировать фактами и числами.
        Если недостаточно информации из файла, используй свой опыт и знания, чтобы раскрыть этот раздел."""

        # добавить 13 в бд, исправить запрос
        question14 = """Представь, что ты главный маркетолог продукта в новой компании, которая только собирается
        открываться и собирается выводить на рынок новый продукт. Информация о продукте содержится в файле, который я тебе прислал.
        Твоя задача - написать текст, в котором будут отражено решение тех болей, проблем, желаний, потребностей из прошлого пункта и указать,
        как наш продукт может закрыть вопрос покупателя. Также необходимо показать в каком объеме достигается решение вопроса.
        Если недостаточно информации из файла, используй свой опыт и знания, чтобы раскрыть этот раздел."""

        question15 = """Представь, что ты главный маркетолог продукта в новой компании, которая
        только собирается открываться и собирается выводить на рынок новый продукт. Информация о продукте содержится в файле, который я тебе прислал.
        Твоя задача - написать текст, в котором будут отражены 3 целевых аудитории нашего продукта, их детальное описание по психографическим признакам,
        их особенности выбора, заинтересованность, мотивация и иные психографические признаки, которые влияют на покупку товара продукта, т.е, те факторы
        которые характеризуют целевую аудиторию и дают ей идентичность относительно другой группы ЦА.
        Если недостаточно информации из файла, используй свой опыт и знания, чтобы раскрыть этот раздел."""

        docx_bytesio_to_md(uploaded_file)
        loader = TextLoader("output_file.txt", encoding='utf-16')
        docs = loader.load()
        # Split documents into chunks
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=200,
            separators=["\n\n", "\n", "(?<=\. )"]
        )
        splitted_texts = splitter.split_documents(docs)
        print(splitted_texts)
        persist_directory = ''
        embedding = OpenAIEmbeddings()
        vectordb = Chroma.from_documents(
            documents=splitted_texts,
            embedding=embedding,
            persist_directory=persist_directory
        )  # создаем БД

        vectordb.persist()  # сохраним БД
        vectordb = Chroma(
            persist_directory=persist_directory,
            embedding_function=embedding
        )
        retriever = vectordb.as_retriever(search_type="mmr", fetch_k=40)
        PROMPT = PromptTemplate.from_template(glob)
        chain_type_kwargs = {"prompt": PROMPT}

        llm = ChatOpenAI(
            model_name="gpt-3.5-turbo-16k",
            temperature=0,
            max_tokens=2048)
        qa_chain = RetrievalQA.from_chain_type(
            llm,
            retriever=retriever,
            chain_type_kwargs=chain_type_kwargs
        )

        async def async_run(chain, question):
            return await loop.run_in_executor(None, chain.run, question)


        loop = asyncio.get_event_loop()

        questions = [question1, question2, question3, question4, question5, question6, question7, question8, question9,
                     question10, question11, question12, question13, question14, question15]
        titles = ["Команда", "Аннотация к проекту", "Какой продукт или услуга будет продаваться",
                  "Какую и чью проблему решает продукт", "Решения используемые в продукте", "Схема организации бизнеса",
                  "Реализуемость", "Технические параметры", "Производственные параметры", "Конкуретные преимущества",
                  "Научно-техническое решение", "Задел", "Характеристика проблемы", "Решение проблемы",
                  "Держатель проблемы"]

        tasks = [loop.create_task(async_run(qa_chain, question)) for question in questions[:9]]
        results = loop.run_until_complete(asyncio.gather(*tasks))

        docs = [MyDocument(result) for result in results]
        splitted_texts = splitter.split_documents(docs)
        vectordb = Chroma.from_documents(
            documents=splitted_texts,
            embedding=embedding,
            persist_directory=persist_directory
        )
        vectordb.persist()

        tasks = [loop.create_task(async_run(qa_chain, question)) for question in questions[9:]]
        results_second_part = loop.run_until_complete(asyncio.gather(*tasks))

        doc = Document()

        for question, result in zip(titles[0:9], results):
            doc.add_heading(question, level=1)
            doc.add_paragraph(result)

        for question, result in zip(titles[10:14], results_second_part):
            doc.add_heading(question, level=1)
            doc.add_paragraph(result)

        return doc


st.set_page_config(page_title='Сократи ВКР')
st.title('Сократи ВКР')

uploaded_file = st.file_uploader('Загрузите текст', type='docx')

with st.form('myform', clear_on_submit=False):
    num_employees = st.radio('Выберите количество сотрудников', options=list(range(1, 6)))
    team_name = st.text_input('Введите название команды')
    project_name = st.text_input('Введите название проекта')

    selected_num_employees = st.session_state.get('selected_num_employees', 1)

    if num_employees != selected_num_employees:
        selected_num_employees = num_employees
        st.session_state['selected_num_employees'] = selected_num_employees

    if st.button('Создать сотрудника'):
        for i in range(selected_num_employees):
            st.subheader(f'Информация о сотруднике {i + 1}')
            employee_name_key = f'employee_name_{i + 1}'
            university_info_key = f'university_info_{i + 1}'
            employee_name = st.text_input(f'ФИО сотрудника {i + 1}', key=employee_name_key)
            university_info = st.text_input(f'ВУЗ, направление, курс сотрудника {i + 1}', key=university_info_key)

    # Проверяем, все ли поля заполнены
    all_fields_filled = all(st.session_state[f'employee_name_{i + 1}'] and st.session_state[f'university_info_{i + 1}']
                            for i in range(selected_num_employees))

    # Проверяем, была ли форма отправлена
    submitted = st.form_submit_button('Сгенерировать документ')

    if submitted and not (uploaded_file and team_name and all_fields_filled):
        unfilled_fields = []

        if not uploaded_file:
            unfilled_fields.append("Файл не загружен.")

        if not team_name:
            unfilled_fields.append("Введите название команды.")

        for i in range(selected_num_employees):
            if not st.session_state[f'employee_name_{i + 1}']:
                unfilled_fields.append(f"ФИО сотрудника {i + 1} не заполнено.")
            if not st.session_state[f'university_info_{i + 1}']:
                unfilled_fields.append(f'ВУЗ, направление, курс сотрудника {i + 1} не заполнено.')

        st.error("Пожалуйста, заполните следующие обязательные поля:\n" + "\n".join(unfilled_fields))
        submitted = False

    submitted = st.form_submit_button('Сгенерировать документ')

    if submitted:
        with st.spinner('Документ обрабатывается...'):
            employees_info = []
            for i in range(selected_num_employees):
                employee_name_key = f'employee_name_{i + 1}'
                university_info_key = f'university_info_{i + 1}'
                employee_name = st.session_state.get(employee_name_key, '')
                university_info = st.session_state.get(university_info_key, '')
                employees_info.append({'employee_name': employee_name, 'university_info': university_info})
            try:
                response = generate_response(uploaded_file, employees_info, project_name, num_employees)
                st.session_state['form_processed'] = True
            except APIError as e:
                st.error(f"Произошла ошибка: {e}, попробуйте подключить VPN")


if st.session_state.get('form_processed'):
    with io.BytesIO() as temp_file:
        response.save(temp_file)
        temp_file.seek(0)


        st.download_button(
            label="Скачать документ",
            data=temp_file,
            file_name="document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
